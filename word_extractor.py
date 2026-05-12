import docx
import io

def extract_text_from_word(file_bytes, filename):
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        text = "WORD DOCUMENT: " + filename + "\n"
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            text += "\n[TABLE]\n"
            for row in table.rows:
                row_values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_values:
                    text += " | ".join(row_values) + "\n"
            text += "[/TABLE]\n"
        return text
    except Exception as e:
        return "Word extraction failed: " + str(e)
